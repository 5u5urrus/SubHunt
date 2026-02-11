#!/usr/bin/env python3
"""
SubHunt - fast subdomain enumeration from massive passive DNS datasets
Author: Vahe Demirkhanyan
"""

import sys
import time
import random
import string
import socket
import argparse
from typing import Any, Iterable, Optional, Set, FrozenSet, List, Tuple, Dict
from concurrent.futures import ThreadPoolExecutor, wait, FIRST_COMPLETED
from urllib.parse import urlparse

import requests

API_URL = "https://ip.thc.org/api/v1/lookup/subdomains"
WAYBACK_CDX_URL = "https://web.archive.org/cdx/search/cdx"
CRTSH_URL = "https://crt.sh/"

CANDIDATE_KEYS = ("domain", "subdomain", "fqdn", "name", "host")

MAX_RETRIES = 6
BASE_BACKOFF_SEC = 0.6
MAX_BACKOFF_SEC = 12.0
TIMEOUT_SEC = 30

TRANSIENT_STATUSES = {408, 425, 429, 500, 502, 503, 504}

# DNS-live filter settings 
DNS_WORKERS = 60
MAX_INFLIGHT = 2500

# wayback/CDX source
WAYBACK_LIMIT = 20000  # keep sane.. dNS-live filter handles "junk", but CDX can be huge


def die(msg: str, code: int = 1) -> None:
    print(msg, file=sys.stderr)
    sys.exit(code)


def extract_domains(obj: Any) -> Iterable[str]:
    if isinstance(obj, str):
        yield obj
        return
    if isinstance(obj, list):
        for item in obj:
            yield from extract_domains(item)
        return
    if isinstance(obj, dict):
        for k in CANDIDATE_KEYS:
            v = obj.get(k)
            if isinstance(v, str) and v.strip():
                yield v.strip()
        for k in ("subdomains", "results", "data", "items"):
            v = obj.get(k)
            if isinstance(v, (list, dict, str)):
                yield from extract_domains(v)
        for v in obj.values():
            if isinstance(v, (list, dict)):
                yield from extract_domains(v)


def find_next_page_state(obj: Any) -> Optional[str]:
    if isinstance(obj, dict):
        for k in ("page_state", "next_page_state", "next", "cursor"):
            v = obj.get(k)
            if isinstance(v, str) and v:
                return v
        for v in obj.values():
            ps = find_next_page_state(v)
            if ps:
                return ps
    elif isinstance(obj, list):
        for v in obj:
            ps = find_next_page_state(v)
            if ps:
                return ps
    return None


def _retry_after_seconds(resp: requests.Response) -> Optional[float]:
    ra = resp.headers.get("Retry-After")
    if not ra:
        return None
    try:
        return float(int(ra.strip()))
    except Exception:
        return None


def _sleep_backoff(attempt: int, forced_min: Optional[float] = None) -> None:
    delay = min(MAX_BACKOFF_SEC, BASE_BACKOFF_SEC * (2 ** attempt))
    delay *= random.uniform(0.85, 1.15)
    if forced_min is not None:
        delay = max(delay, forced_min)
    time.sleep(delay)


def post_lookup(session: requests.Session, domain: str, limit: int, page_state: str) -> Any:
    payload = {"domain": domain, "limit": limit, "page_state": page_state}
    headers = {"Accept": "application/json", "User-Agent": "subhunt/1.5"}

    last_err: Optional[str] = None

    for attempt in range(MAX_RETRIES):
        try:
            r = session.post(API_URL, json=payload, headers=headers, timeout=TIMEOUT_SEC)
        except requests.RequestException as e:
            last_err = f"Request error: {e}"
            _sleep_backoff(attempt)
            continue

        if 200 <= r.status_code < 300:
            try:
                return r.json()
            except Exception:
                snippet = r.text[:200].replace("\n", "\\n")
                die(f"Unexpected response (not JSON). First 200 chars: {snippet}", 2)

        if r.status_code in TRANSIENT_STATUSES or (500 <= r.status_code <= 599):
            ra = _retry_after_seconds(r) if r.status_code == 429 else None
            last_err = f"HTTP {r.status_code}"
            _sleep_backoff(attempt, forced_min=ra)
            continue

        snippet = (r.text or "")[:200].replace("\n", "\\n")
        die(f"HTTP {r.status_code}. First 200 chars: {snippet}", 2)

    die(f"Failed after {MAX_RETRIES} retries. Last error: {last_err or 'unknown'}", 2)


def _rand_label(n: int = 18) -> str:
    alphabet = string.ascii_lowercase + string.digits
    return "".join(random.choice(alphabet) for _ in range(n))


def resolve_host_ips(host: str) -> FrozenSet[str]:
    try:
        infos = socket.getaddrinfo(host, None, type=socket.SOCK_STREAM)
    except socket.gaierror:
        return frozenset()
    except Exception:
        return frozenset()

    ips: Set[str] = set()
    for _family, _socktype, _proto, _canonname, sockaddr in infos:
        if isinstance(sockaddr, tuple) and sockaddr:
            ip = sockaddr[0]
            if isinstance(ip, str) and ip:
                ips.add(ip)
    return frozenset(ips)


def detect_wildcard_signature(domain: str) -> Optional[FrozenSet[str]]:
    ipsets: List[FrozenSet[str]] = []
    for _ in range(3):
        host = f"{_rand_label()}.{domain}"
        ipset = resolve_host_ips(host)
        if ipset:
            ipsets.append(ipset)

    if len(ipsets) < 2:
        return None

    counts: Dict[FrozenSet[str], int] = {}
    for s in ipsets:
        counts[s] = counts.get(s, 0) + 1

    best_set, best_count = max(counts.items(), key=lambda kv: kv[1])
    if best_count >= 2:
        return best_set
    return None


def _get_json_with_retries(session: requests.Session, url: str, params: dict) -> Any:
    headers = {"Accept": "application/json", "User-Agent": "subhunt/1.5"}
    last_err: Optional[str] = None

    for attempt in range(MAX_RETRIES):
        try:
            r = session.get(url, params=params, headers=headers, timeout=TIMEOUT_SEC)
        except requests.RequestException as e:
            last_err = f"Request error: {e}"
            _sleep_backoff(attempt)
            continue

        if 200 <= r.status_code < 300:
            try:
                return r.json()
            except Exception:
                last_err = "Non-JSON response"
                _sleep_backoff(attempt)
                continue

        if r.status_code in TRANSIENT_STATUSES or (500 <= r.status_code <= 599):
            ra = _retry_after_seconds(r) if r.status_code == 429 else None
            last_err = f"HTTP {r.status_code}"
            _sleep_backoff(attempt, forced_min=ra)
            continue

        last_err = f"HTTP {r.status_code}"
        break

    return None


def fetch_wayback_candidates(session: requests.Session, domain: str) -> Iterable[str]:
    params = {
        "url": f"*.{domain}/*",
        "output": "json",
        "fl": "original",
        "collapse": "urlkey",
        "limit": str(WAYBACK_LIMIT),
    }

    data = _get_json_with_retries(session, WAYBACK_CDX_URL, params=params)
    if not isinstance(data, list) or len(data) < 2:
        return []

    out: List[str] = []
    for row in data[1:]:
        try:
            if not isinstance(row, list) or not row:
                continue
            raw_url = row[0]
            if not isinstance(raw_url, str) or not raw_url:
                continue
            host = urlparse(raw_url).hostname
            if host:
                out.append(host.strip().lower().rstrip("."))
        except Exception:
            continue

    return out


def fetch_crtsh_candidates(session: requests.Session, domain: str) -> Iterable[str]:
    params = {"q": f"%.{domain}", "output": "json"}

    data = _get_json_with_retries(session, CRTSH_URL, params=params)
    if not isinstance(data, list) or not data:
        return []

    out: List[str] = []
    for entry in data:
        try:
            if not isinstance(entry, dict):
                continue
            name_value = entry.get("name_value")
            if not isinstance(name_value, str) or not name_value:
                continue
            for host in name_value.splitlines():
                host = host.strip().lower().rstrip(".")
                if host:
                    if host.startswith("*."):
                        host = host[2:]
                    out.append(host)
        except Exception:
            continue

    return out


def _write_docx(path: str, domain: str, rows: List[Tuple[str, str]]) -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception:
        print("Warning: python-docx not installed; cannot write DOCX output.", file=sys.stderr)
        return

    if not path.lower().endswith(".docx"):
        path = path + ".docx"

    doc = Document()

    title = doc.add_paragraph(f"SubHunt results for: {domain}")
    for run in title.runs:
        run.font.name = "Calibri Light"
        run.font.size = Pt(9)

    doc.add_paragraph("")

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Subdomain"
    hdr[1].text = "IPs"

    def style_cell_text(cell):
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Calibri Light"
                run.font.size = Pt(9)

    style_cell_text(hdr[0])
    style_cell_text(hdr[1])

    for i, (host, ips) in enumerate(rows, start=1):
        c0, c1 = table.rows[i].cells
        c0.text = host
        c1.text = ips
        style_cell_text(c0)
        style_cell_text(c1)

    try:
        doc.save(path)
    except Exception as e:
        print(f"Warning: Could not write DOCX file '{path}': {e}", file=sys.stderr)


def _parse_args(argv: List[str]) -> argparse.Namespace:
    p = argparse.ArgumentParser(
        prog=(argv[0] if argv else "subhunt.py"),
        add_help=True,
        description="SubHunt - fast subdomain enumeration with DNS-live filtering.",
    )
    p.add_argument("domain", help="Target domain, e.g. example.com")
    p.add_argument("output", nargs="?", default=None, help="Optional output DOCX path")
    p.add_argument(
        "--full",
        action="store_true",
        help="Also use Wayback CDX + crt.sh (slower, more coverage)",
    )
    return p.parse_args(argv[1:])


def main(argv: Optional[List[str]] = None) -> None:
    if argv is None:
        argv = sys.argv

    args = _parse_args(argv)

    domain = (args.domain or "").strip().lower().rstrip(".")
    if not domain or "/" in domain or " " in domain:
        die("Invalid domain")

    out_docx = (args.output.strip() if isinstance(args.output, str) else None) if args.output else None

    session = requests.Session()
    wildcard_sig = detect_wildcard_signature(domain)

    def is_in_scope(host: str) -> bool:
        return host == domain or host.endswith("." + domain)

    found_rows: List[Tuple[str, str]] = []
    found_set: Set[str] = set()

    def resolve_pair(h: str) -> Tuple[str, FrozenSet[str]]:
        return (h, resolve_host_ips(h))

    with ThreadPoolExecutor(max_workers=DNS_WORKERS) as ex:
        inflight = set()

        def submit_candidate(host: str):
            return ex.submit(resolve_pair, host)

        def record_and_print(host: str, ipset: FrozenSet[str]) -> None:
            if host in found_set:
                return
            found_set.add(host)

            print(host, flush=True)

            if out_docx is None:
                return

            ips = ", ".join(sorted(ipset)) if ipset else ""
            found_rows.append((host, ips))

        def drain_some(block: bool) -> None:
            nonlocal inflight
            if not inflight:
                return

            timeout = None if block else 0
            done, pending = wait(inflight, timeout=timeout, return_when=FIRST_COMPLETED)

            if not done:
                return

            inflight = pending
            for f in done:
                try:
                    host, ipset = f.result()
                except Exception:
                    continue
                if not ipset:
                    continue
                if wildcard_sig is not None and ipset == wildcard_sig:
                    continue
                record_and_print(host, ipset)

        def enqueue(host: str, seen: Set[str]) -> None:
            host = host.strip().lower().rstrip(".")
            if not host or not is_in_scope(host):
                return
            if host in seen:
                return
            seen.add(host)

            inflight.add(submit_candidate(host))
            while len(inflight) >= MAX_INFLIGHT:
                drain_some(block=True)

        seen: Set[str] = set()

        # -------- Passive source #1: THC (default) --------------------
        page_state = ""
        limit = 500

        while True:
            obj = post_lookup(session, domain, limit, page_state)
            for s in extract_domains(obj):
                enqueue(s, seen)

            next_state = find_next_page_state(obj)
            if not next_state or next_state == page_state:
                break

            page_state = next_state
            time.sleep(0.15)

            for _ in range(4):
                if not inflight:
                    break
                drain_some(block=False)

        # -------- Optional extra sources (only with --full) -------
        if args.full:
            for s in fetch_wayback_candidates(session, domain):
                enqueue(s, seen)
                if inflight and random.random() < 0.08:
                    drain_some(block=False)

            for s in fetch_crtsh_candidates(session, domain):
                enqueue(s, seen)
                if inflight and random.random() < 0.08:
                    drain_some(block=False)

        while inflight:
            drain_some(block=True)

    if out_docx is not None:
        found_rows.sort(key=lambda x: x[0])
        _write_docx(out_docx, domain, found_rows)


if __name__ == "__main__":
    main()
